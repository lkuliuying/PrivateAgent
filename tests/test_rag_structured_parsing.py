from __future__ import annotations

import sys
from itertools import pairwise
from types import SimpleNamespace

import pytest

from personal_assistant.core.rag import (
    estimate_token_count,
    parse_document_blocks,
    split_document_blocks,
)


def test_markdown_blocks_retain_exact_offsets_lines_and_heading_path(tmp_path):
    raw = (
        "# Root\n\n"
        "intro line\nsecond line\n\n"
        "## Child\n"
        "child body one\nchild body two\n"
    )
    source = tmp_path / "structured.md"
    source.write_text(raw, encoding="utf-8")

    blocks = parse_document_blocks(str(source))

    assert [block.source_kind for block in blocks] == [
        "markdown_heading",
        "markdown_block",
        "markdown_heading",
        "markdown_block",
    ]
    assert blocks[0].heading_path == ("Root",)
    assert blocks[2].heading_path == ("Root", "Child")
    assert blocks[3].heading_path == ("Root", "Child")
    assert blocks[3].line_start == 7
    assert blocks[3].line_end == 8
    for block in blocks:
        assert raw[block.char_start : block.char_end] == block.content

    chunks = split_document_blocks(blocks, size=12, overlap=3)
    assert chunks
    for chunk in chunks:
        assert raw[chunk.char_start : chunk.char_end] == chunk.content
        assert chunk.parser_version == "markdown:v2"
        assert chunk.line_start is not None
        assert chunk.line_end is not None
        assert chunk.line_start <= chunk.line_end


def test_markdown_fenced_code_is_one_semantic_block(tmp_path):
    raw = (
        "# Root\n\n"
        "```python\n"
        "# this is code, not a heading\n"
        "\n"
        "print('ready')\n"
        "```\n\n"
        "## Child\n"
        "after fence\n"
    )
    source = tmp_path / "fenced.md"
    source.write_text(raw, encoding="utf-8")

    blocks = parse_document_blocks(str(source))

    assert [block.source_kind for block in blocks] == [
        "markdown_heading",
        "markdown_code_fence",
        "markdown_heading",
        "markdown_block",
    ]
    fenced = blocks[1]
    assert fenced.heading_path == ("Root",)
    assert fenced.parser_version == "markdown:v2"
    assert "# this is code, not a heading" in fenced.content
    assert raw[fenced.char_start : fenced.char_end] == fenced.content
    assert blocks[2].heading_path == ("Root", "Child")

    chunks = split_document_blocks(blocks, size=24, overlap=4)
    fenced_chunks = [
        chunk for chunk in chunks if chunk.source_kind == "markdown_code_fence"
    ]
    assert fenced_chunks
    assert all(chunk.heading_path == ("Root",) for chunk in fenced_chunks)
    assert all(chunk.parser_version == "markdown:v2" for chunk in fenced_chunks)


def test_docx_tables_preserve_document_order_and_heading_scope(tmp_path):
    from docx import Document

    source = tmp_path / "structured.docx"
    document = Document()
    document.add_heading("Root", level=1)
    document.add_paragraph("before table")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "name"
    table.cell(0, 1).text = "value"
    table.cell(1, 0).text = "alpha|beta"
    table.cell(1, 1).text = "C:\\temp"
    document.add_paragraph("after table")
    document.save(source)

    blocks = parse_document_blocks(str(source))

    assert [block.source_kind for block in blocks] == [
        "docx_paragraph",
        "docx_paragraph",
        "docx_table",
        "docx_paragraph",
    ]
    assert [block.content for block in blocks[:2]] == ["Root", "before table"]
    table_block = blocks[2]
    assert table_block.content == (
        "| name | value |\n| alpha\\|beta | C:\\\\temp |"
    )
    assert table_block.heading_path == ("Root",)
    assert blocks[3].content == "after table"
    assert all(block.parser_version == "python-docx:v2" for block in blocks)
    assert all(
        left.char_end < right.char_start for left, right in pairwise(blocks)
    )


def test_pdf_chunks_never_cross_pages_and_retain_page_coordinates(monkeypatch):
    page_texts = [
        "Page one title\n" + "alpha beta gamma " * 6,
        "Page two title\n" + "delta epsilon zeta " * 6,
    ]

    class FakePage:
        def __init__(self, content: str) -> None:
            self.content = content

        def extract_text(self) -> str:
            return self.content

    class FakeReader:
        def __init__(self, path: str) -> None:
            del path
            self.pages = [FakePage(content) for content in page_texts]

    monkeypatch.setitem(
        sys.modules,
        "pypdf",
        SimpleNamespace(PdfReader=FakeReader),
    )

    blocks = parse_document_blocks("synthetic.pdf")
    chunks = split_document_blocks(blocks, size=32, overlap=5)
    extracted = "\n".join(page_texts)

    assert [block.page_start for block in blocks] == [1, 2]
    assert all(block.page_start == block.page_end for block in blocks)
    assert {chunk.page_start for chunk in chunks} == {1, 2}
    for chunk in chunks:
        assert chunk.page_start == chunk.page_end
        assert extracted[chunk.char_start : chunk.char_end] == chunk.content
        assert chunk.parser_version == "pypdf:v1"


def test_python_code_blocks_follow_top_level_symbols_and_decorators(tmp_path):
    raw = (
        '"""module docs"""\nVALUE = 1\n\n'
        "@trace\n"
        "def first(value):\n"
        "    return value\n\n"
        "class Worker:\n"
        "    def run(self):\n"
        "        return VALUE\n"
    )
    source = tmp_path / "worker.py"
    source.write_text(raw, encoding="utf-8")

    blocks = parse_document_blocks(str(source))

    assert [block.source_kind for block in blocks] == [
        "code_module",
        "code_symbol",
        "code_symbol",
    ]
    assert blocks[1].heading_path == ("worker.py", "function first")
    assert blocks[1].line_start == 4
    assert blocks[2].heading_path == ("worker.py", "class Worker")
    assert blocks[2].line_start == 8
    assert all(block.parser_version == "python-ast:v1" for block in blocks)
    for block in blocks:
        assert raw[block.char_start : block.char_end] == block.content


def test_generic_code_and_plain_text_use_structure_before_windowing(tmp_path):
    typescript = (
        "const prefix = 'x'\n\n"
        "export class Greeter {\n  hello() { return prefix }\n}\n\n"
        "export async function load() { return 1 }\n"
    )
    code_source = tmp_path / "greeter.ts"
    code_source.write_text(typescript, encoding="utf-8")
    code_blocks = parse_document_blocks(str(code_source))
    assert [block.source_kind for block in code_blocks] == [
        "code_module",
        "code_symbol",
        "code_symbol",
    ]
    assert code_blocks[1].heading_path[-1] == "class Greeter"
    assert code_blocks[2].heading_path[-1] == "function load"

    plain = "first paragraph\ncontinues\n\nsecond paragraph\n"
    text_source = tmp_path / "notes.txt"
    text_source.write_text(plain, encoding="utf-8")
    text_blocks = parse_document_blocks(str(text_source))
    assert [block.content for block in text_blocks] == [
        "first paragraph\ncontinues",
        "second paragraph",
    ]
    assert [block.line_start for block in text_blocks] == [1, 4]
    assert all(block.source_kind == "text_paragraph" for block in text_blocks)


def test_invalid_python_records_a_bounded_parse_error(tmp_path):
    source = tmp_path / "broken.py"
    source.write_text("def broken(:\n    pass\n", encoding="utf-8")

    with pytest.raises(ValueError, match=r"Python 代码解析失败（line 1）"):
        parse_document_blocks(str(source))


def test_token_estimate_is_nonzero_and_conservative_for_cjk():
    assert estimate_token_count("") == 0
    assert estimate_token_count("abcdef") == 2
    assert estimate_token_count("中文abc") == 3


@pytest.mark.parametrize(
    ("size", "overlap"),
    [(0, 0), (10, -1), (10, 10), (10, 11)],
)
def test_structured_split_rejects_invalid_window(size, overlap):
    with pytest.raises(ValueError, match="overlap"):
        split_document_blocks([], size=size, overlap=overlap)
